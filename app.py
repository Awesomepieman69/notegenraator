from flask import Flask, render_template, request, send_file
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from PyPDF2 import PdfReader
import re
import google.generativeai as genai

load_dotenv() # Load environment variables from .env file

app = Flask(__name__)

# Configure Google AI Client (looks for GOOGLE_API_KEY in .env)
try:
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
except Exception as e:
    print(f"Error configuring Google AI SDK: {e}. Make sure GOOGLE_API_KEY is set in .env")
    # Optionally exit or handle appropriately if key is essential at startup

# --- Helper Functions ---

# Updated for Google Gemini
def get_completion(prompt, model_name="gemini-2.5-flash-preview-04-17"):
    """Calls the Google Gemini API to get a completion."""
    try:
        model = genai.GenerativeModel(model_name)
        # Basic safety settings - adjust as needed
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        ]
        # Generation config - temperature=0 for deterministic output
        generation_config = genai.types.GenerationConfig(temperature=0)

        response = model.generate_content(
            prompt, 
            generation_config=generation_config,
            safety_settings=safety_settings
        )
        
        # Check for valid response and text part
        if response.parts:
            return response.text
        else:
            # Handle potential blocks or empty responses
            print(f"Warning: Gemini response finished reason: {response.prompt_feedback.block_reason}")
            print(f"Full response feedback: {response.prompt_feedback}")
            return None # Indicate failure or blocked content

    except Exception as e:
        print(f"Error calling Google Gemini API: {e}")
        return None

def create_styled_docx(content):
    """Creates a Word document with specific styling for headings marked with ##."""
    document = Document()
    lines = content.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Updated heading detection based on ## marker
        if line.startswith('## '):
            heading_text = line[3:].strip() # Remove '## ' prefix
            paragraph = document.add_paragraph()
            run = paragraph.add_run(heading_text)
            run.font.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text
            
            # Add blue background shading to the paragraph
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4A86E8') # Approximate blue color
            shading_elm.set(qn('w:val'), 'clear')
            paragraph.paragraph_format.element.get_or_add_pPr().append(shading_elm)
            paragraph.paragraph_format.space_after = Pt(10)
        # Revised bullet point detection and formatting based on first colon
        elif line.startswith(('- ', '* ', '• ')):
            content = line.lstrip('-*• ').strip() # Remove bullet marker
            paragraph = document.add_paragraph(style='List Bullet')
            
            # Find the first colon
            colon_index = content.find(':')
            
            if colon_index != -1: # If colon exists
                term_part = content[:colon_index]
                description_part = content[colon_index+1:] # Text after colon
                
                # Clean term: remove leading/trailing spaces and asterisks
                cleaned_term = term_part.strip().strip('*').strip()
                
                # Clean description: remove leading/trailing spaces and asterisks
                cleaned_description = description_part.strip().strip('*').strip()
                
                # Add cleaned term as bold
                run = paragraph.add_run(cleaned_term)
                run.bold = True
                
                # Add colon and space as bold
                run = paragraph.add_run(': ')
                run.bold = True
                
                # Add cleaned description as normal text
                paragraph.add_run(cleaned_description)
                
            else:
                # No colon found, add the whole content normally
                # Clean content just in case
                cleaned_content = content.strip().strip('*').strip()
                paragraph.add_run(cleaned_content)
        else:
            # Regular paragraph
            document.add_paragraph(line)

    return document

# --- Helper function to parse numbered Q&A response ---
def parse_numbered_answers(response_text):
    """
    Returns a dict {index:int -> answer:str} for lines like 'A3: …'.
    Handles multi‑line answers.
    """
    pattern = re.compile(r"^A(\d+):\s*(.*?)(?=\nA\d+:|\Z)", re.MULTILINE | re.DOTALL)
    # Use a dictionary comprehension for cleaner creation
    return {int(i): ans.strip() for i, ans in pattern.findall(response_text or "")} # Handle None response_text

# --- Routes ---

@app.route('/')
def index():
    """Renders the homepage."""
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_document():
    """Handles file upload, calls OpenAI based on selected action, and returns DOCX."""
    action = request.form.get('action', 'generate_notes')
    output_filename = "generated_notes.docx" if action == 'generate_notes' else "answered_questions.docx"

    if 'files[]' not in request.files:
        return "No files selected", 400

    files = request.files.getlist('files[]')

    prompt_template = "" # Initialize prompt template
    original_paragraphs = [] # Initialize for QA action

    # --- Action Specific Logic --- 
    if action == 'answer_questions':
        if len(files) != 1:
            return "Please upload exactly one document for paragraph-by-paragraph answering.", 400
        file = files[0]
        if not file.filename.lower().endswith('.docx'):
            return "Paragraph answering currently only supports .docx files.", 400
        if file.filename == '':
             return "No valid file selected", 400

        # --- Read original docx structure ---
        try:
            original_doc = Document(file.stream)
            for idx, para in enumerate(original_doc.paragraphs): # Add enumerate for index
                if para.text.strip():
                    # Store index along with text and style
                    original_paragraphs.append({
                        "index": idx + 1, # Store 1-based index
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })
            if not original_paragraphs:
                 return "The document contains no text paragraphs to process.", 400
        except Exception as e:
            print(f"Error reading docx file {file.filename}: {e}")
            return f"Error reading Word document: {e}", 500

        # --- Construct the Q&A prompt (General Knowledge for Every Paragraph) ---
        # Prefix questions with running index (Q1:, Q2:, ...)
        paragraphs_for_prompt = "\n".join(
            [f"Q{p['index']}: {p['text']}" for p in original_paragraphs]
        )
        # Update prompt instructions for numbered Q/A format
        prompt_template = f"""You are a helpful AI assistant.
For each numbered question below (Q1, Q2, etc.), provide a concise answer or response based on your general knowledge on the **very next line**.
Use the format **A<same number>: [Your Response/Answer]** for each answer. Treat each question independently.

Example (follow this exactly):
Q1: What is the capital of France?
A1: Paris
Q2: Summarize the main idea of this sentence.
A2: The sentence asks for a summary of its own main idea.

Questions from document:
{paragraphs_for_prompt}
"""

    elif action == 'generate_notes':
        # --- Read files for Note Generation ---
        full_document_text = ""
        for file in files:
            if file.filename == '': continue
            if file:
                extracted_content = ""
                try:
                    filename_lower = file.filename.lower()
                    if filename_lower.endswith('.docx'):
                        doc = Document(file.stream)
                        extracted_content = "\n".join([para.text for para in doc.paragraphs])
                    elif filename_lower.endswith('.pdf'):
                        reader = PdfReader(file.stream)
                        text_parts = []
                        for page in reader.pages:
                            text_parts.append(page.extract_text() or "")
                        extracted_content = "\n".join(text_parts)
                    else:
                        try: extracted_content = file.read().decode('utf-8')
                        except UnicodeDecodeError:
                            print(f"Warning: Could not decode file {file.filename} as UTF-8. Skipping.")
                            continue
                    if extracted_content.strip():
                        full_document_text += extracted_content.strip() + "\n\n---\n\n"
                    else: print(f"Warning: No text content extracted from file {file.filename}.")
                except Exception as e: print(f"Error processing file {file.filename}: {e}")
        if not full_document_text: return "No valid document content found", 400
        
        # --- Construct Note Generation Prompt ---
        prompt_template = f"""You are an expert note-taker. Your task is to generate comprehensive, accurate, and well-structured notes from the provided transcript of a lecture, video, or similar source material. Follow the universal outline specified below precisely. **Keep the overall notes as brief and high-level as possible while still covering the main points according to the outline.**

**Instructions & Guidelines:**
1.  **Accuracy is Paramount:** Only include information that is *explicitly stated* in the transcript. Do NOT add external knowledge, assumptions, or interpretations.
2.  **Thoroughness:** Extract all relevant details for each section and bullet point. Ensure explanations are **very concise** (typically **1 brief sentence or phrase** per bullet point) and focused *only* on the transcript's content.
3.  **Structure:** Start each section heading on a new line, prefixed with `## `. List key items within sections using bullet points starting with `- `.
4.  **Handling Missing Information:** If the transcript lacks information for a specific section (e.g., no examples were mentioned), skip that section *entirely*. Do not include the heading or any placeholder text for missing sections.
5.  **Clarity:** Use clear and direct language.

**Outline to Follow:**
⸻
## Section 1: Introduction and Overview
Provide a brief summary of the lecture's main themes and objectives, using only information present in the transcript. Aim for 1-2 complete sentences setting the context.

## Section 2: Main Concepts and Key Points
- State the first key idea or concept found in the transcript. Explain its significance or provide key details using 1-2 sentences based *only* on the transcript.
- List the next important idea found, followed by 1-2 sentences of elaboration derived *only* from the transcript.
- Continue listing all distinct key concepts/points mentioned, each with 1-2 sentences of transcript-based explanation.

## Section 3: Definitions and Important Terms
- Identify a significant term defined or explained in the transcript. Provide its definition or explanation using 1-2 sentences based *only* on the transcript.
- Continue listing all distinct terms defined/explained, each with 1-2 sentences of transcript-based explanation.

## Section 4: Examples, Illustrations, or Case Studies
- Describe the first example, illustration, or case study mentioned in the transcript. Explain its relevance or purpose using 1-2 sentences based *only* on the transcript.
- Continue listing all distinct examples/illustrations/case studies mentioned, each with 1-2 sentences of transcript-based explanation.

## Section 5: Summary and Key Takeaways
DO NOT EXCEED more than 4 pages in length for document.
Summarize the main conclusions or key take-home messages explicitly stated at the end of the transcript. Aim for 1-2 concise but comprehensive sentences.
⸻
TRANSCRIPT:
{full_document_text}
"""

    else:
        return "Invalid action specified", 400
    # --- End Action Specific Logic ---

    # --- Token Counting (using Google SDK) ---
    MODEL_NAME = "gemini-2.5-flash-preview-04-17" # User specified model
    # Using a very high limit due to Gemini's large context window
    MAX_TOKENS_ALLOWED = 900000 
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        num_tokens = model.count_tokens(prompt_template).total_tokens
        print(f"Estimated prompt tokens (Gemini): {num_tokens}")
    except Exception as e:
        print(f"Error counting tokens with Google SDK: {e}")
        return "Error estimating request size", 500

    if num_tokens > MAX_TOKENS_ALLOWED:
        print(f"Error: Exceeded Gemini token limit. Estimated tokens: {num_tokens}, Limit: {MAX_TOKENS_ALLOWED}")
        error_context = "document paragraphs" if action == 'answer_questions' else "document(s)"
        return f"The combined text for {error_context} is too long for the Gemini model (estimated {num_tokens} tokens, limit ~{MAX_TOKENS_ALLOWED}). Please use shorter/fewer documents.", 413
    # --- End Token Counting ---

    # --- Call AI (Applies to both actions) ---
    generated_content = get_completion(prompt_template, model_name=MODEL_NAME)
    if generated_content is None: # Changed check for None
        return "Error generating response from AI, or content was blocked.", 500
    # --- End AI Call ---

    # --- Document Generation --- 
    file_stream = io.BytesIO()
    if action == 'answer_questions':
        # --- Reconstruct docx with inserted GENERAL KNOWLEDGE answers for every paragraph ---
        try:
            # Use the new parser for numbered answers
            answer_map = parse_numbered_answers(generated_content)
            if not answer_map:
                 print("Warning: Could not parse any numbered answers (A#: ...) from AI response.")

            new_doc = Document() # Create a new document to build
            for para_info in original_paragraphs:
                # Get info for the current original paragraph
                text = para_info["text"].strip() # Strip leading/trailing whitespace
                style_name = para_info["style"] # Get original style name
                para_index = para_info["index"] # Get the stored index

                # 1) If it's a Heading style, preserve it exactly
                #    (Ensure style_name is not None before checking startswith)
                if style_name and style_name.startswith("Heading"):
                    # Attempt to apply the original style if it exists in the new document
                    try:
                        new_doc.add_paragraph(text, style=new_doc.styles[style_name])
                    except KeyError:
                        print(f"Warning: Heading style '{style_name}' not found in new document. Adding as plain text.")
                        new_doc.add_paragraph(text) # Add as plain text if style missing

                # 2) Otherwise it's a "question" — force no indent and bold
                else:
                    q_para = new_doc.add_paragraph()
                    q_para.paragraph_format.left_indent       = Inches(0)
                    q_para.paragraph_format.first_line_indent = Inches(0)
                    q_para.paragraph_format.space_after       = Pt(4)
                    run = q_para.add_run(text)
                    run.bold = True

                # 3) Then insert the answer (if any), indented under its question/heading
                answer_text = answer_map.get(para_index)
                if answer_text:
                    a_para = new_doc.add_paragraph()
                    # Indent the answer paragraph
                    a_para.paragraph_format.left_indent = Inches(0.25)
                    # Add space after the answer
                    a_para.paragraph_format.space_after = Pt(6)
                    # Add the answer text (not bold)
                    a_para.add_run(answer_text)
                # No need for an else block here, if no answer, nothing is added.

            new_doc.save(file_stream)
        except Exception as e:
             print(f"Error reconstructing document with answers: {e}")
             return "Error creating document with answers", 500
        # --- End Reconstruct docx ---
    
    elif action == 'generate_notes':
        # Use existing styled docx creation for notes
        try:
            document = create_styled_docx(generated_content)
            document.save(file_stream)
        except Exception as e:
             print(f"Error creating styled notes document: {e}")
             return "Error creating notes document", 500
    
    else: # Should not happen if action validation is correct
        return "Internal error: Unhandled action for document generation", 500
        
    file_stream.seek(0)
    # --- End Document Generation ---

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=output_filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=True) 